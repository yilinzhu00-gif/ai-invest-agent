"""Immutable researcher-edited brief snapshots and deterministic exports."""

from __future__ import annotations

import hashlib
import html
import json
from enum import Enum
from io import BytesIO
from typing import Literal

from docx import Document
from docx.shared import Pt
from pydantic import BaseModel, ConfigDict, Field, field_validator, model_validator
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

BRIEF_SECTION_TITLES = (
    "已证实的交易事实",
    "公告后的市场反应",
    "可能的影响机制",
    "正面因素",
    "风险和不确定性",
)
DEFAULT_RISK_DISCLAIMER = "本简报仅基于所列来源和数据日期整理，不构成投资建议。"


class BriefExportFormat(str, Enum):
    MARKDOWN = "markdown"
    PDF = "pdf"
    DOCX = "docx"


class BriefCitation(BaseModel):
    model_config = ConfigDict(extra="forbid")

    evidence_id: str = Field(min_length=1, max_length=256)
    filename: str = Field(min_length=1, max_length=512)
    document_version: int = Field(ge=1)
    page_number: int = Field(ge=1)
    block_id: str = Field(min_length=1, max_length=128)


class BriefClaim(BaseModel):
    model_config = ConfigDict(extra="forbid")

    text: str = Field(min_length=1, max_length=10_000)
    citations: list[BriefCitation] = Field(min_length=1, max_length=8)


class BriefSection(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str = Field(min_length=1, max_length=64)
    claims: list[BriefClaim] = Field(default_factory=list, max_length=6)


class ResearchBriefContent(BaseModel):
    """One researcher-owned snapshot; all export formats render this object."""

    model_config = ConfigDict(extra="forbid")

    title: str = Field(default="研究简报", min_length=1, max_length=200)
    summary: str = Field(min_length=1, max_length=10_000)
    data_date: str = Field(pattern=r"^\d{4}-\d{2}-\d{2}$")
    sections: list[BriefSection] = Field(min_length=5, max_length=5)
    missing_information: list[str] = Field(default_factory=list, max_length=16)
    confidence: Literal["high", "medium", "low"]
    confidence_rationale: str = Field(min_length=1, max_length=2_000)
    risk_disclaimer: str = Field(
        default=DEFAULT_RISK_DISCLAIMER, min_length=1, max_length=2_000
    )

    @field_validator("data_date")
    @classmethod
    def require_real_calendar_date(cls, value: str) -> str:
        from datetime import date

        date.fromisoformat(value)
        return value

    @model_validator(mode="after")
    def require_fixed_section_order(self) -> ResearchBriefContent:
        if tuple(section.title for section in self.sections) != BRIEF_SECTION_TITLES:
            raise ValueError("brief sections must use the fixed conclusion order")
        return self


class SaveResearchBriefRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    content: ResearchBriefContent


class ResearchBriefVersion(BaseModel):
    model_config = ConfigDict(extra="forbid")

    version: int = Field(ge=1)
    content: ResearchBriefContent
    source: Literal["researcher"] = "researcher"
    content_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")


class DecideResearchBriefRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    decision: Literal["accept", "reject"]


def content_sha256(content: ResearchBriefContent) -> str:
    canonical = json.dumps(
        content.model_dump(mode="json"), ensure_ascii=False, sort_keys=True, separators=(",", ":")
    )
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def citation_label(citation: BriefCitation) -> str:
    return (
        f"{citation.filename} · v{citation.document_version} · "
        f"第 {citation.page_number} 页 · 块 {citation.block_id}"
    )


def render_markdown(version: ResearchBriefVersion) -> str:
    content = version.content
    lines = [
        f"# {content.title}",
        "",
        f"研究简报版本：v{version.version}",
        f"数据日期：{content.data_date}",
        f"内容校验指纹：{version.content_sha256}",
        "",
        "## 摘要",
        content.summary,
    ]
    for section in content.sections:
        lines.extend(("", f"## {section.title}"))
        if not section.claims:
            lines.append("- 尚缺少可引用证据。")
        for claim in section.claims:
            lines.append(f"- {claim.text}")
            lines.extend(f"  - 来源：{citation_label(citation)}" for citation in claim.citations)
    lines.extend(("", "## 尚缺少的信息"))
    if content.missing_information:
        lines.extend(f"- {item}" for item in content.missing_information)
    else:
        lines.append("- 无。")
    lines.extend(
        (
            "",
            "## 结论置信度",
            f"- {content.confidence}：{content.confidence_rationale}",
            "",
            "## 风险声明",
            content.risk_disclaimer,
            "",
        )
    )
    return "\n".join(lines)


def _document_lines(version: ResearchBriefVersion) -> list[tuple[str, str]]:
    content = version.content
    lines = [
        ("title", content.title),
        ("meta", f"研究简报版本：v{version.version}"),
        ("meta", f"数据日期：{content.data_date}"),
        ("meta", f"内容校验指纹：{version.content_sha256}"),
        ("heading", "摘要"),
        ("body", content.summary),
    ]
    for section in content.sections:
        lines.append(("heading", section.title))
        if not section.claims:
            lines.append(("body", "尚缺少可引用证据。"))
        for claim in section.claims:
            lines.append(("body", claim.text))
            lines.extend(("source", f"来源：{citation_label(citation)}") for citation in claim.citations)
    lines.append(("heading", "尚缺少的信息"))
    if content.missing_information:
        lines.extend(("body", item) for item in content.missing_information)
    else:
        lines.append(("body", "无。"))
    lines.extend(
        (
            ("heading", "结论置信度"),
            ("body", f"{content.confidence}：{content.confidence_rationale}"),
            ("heading", "风险声明"),
            ("body", content.risk_disclaimer),
        )
    )
    return lines


def render_docx(version: ResearchBriefVersion) -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    for kind, text in _document_lines(version):
        if kind == "title":
            paragraph = document.add_heading(text, level=0)
        elif kind == "heading":
            paragraph = document.add_heading(text, level=1)
        else:
            paragraph = document.add_paragraph(text)
            if kind == "source":
                paragraph.style = document.styles["Caption"]
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(version: ResearchBriefVersion) -> bytes:
    """Render the same content snapshot with a built-in Chinese CID font."""
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    output = BytesIO()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("BriefTitle", parent=styles["Title"], fontName="STSong-Light", fontSize=18))
    styles.add(ParagraphStyle("BriefHeading", parent=styles["Heading2"], fontName="STSong-Light", fontSize=13))
    styles.add(ParagraphStyle("BriefBody", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=16))
    styles.add(ParagraphStyle("BriefMeta", parent=styles["BodyText"], fontName="STSong-Light", fontSize=8.5, leading=12))
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story = []
    for kind, text in _document_lines(version):
        style = {"title": "BriefTitle", "heading": "BriefHeading", "meta": "BriefMeta"}.get(
            kind, "BriefBody"
        )
        prefix = "• " if kind == "source" else ""
        story.append(Paragraph(html.escape(prefix + text), styles[style]))
        story.append(Spacer(1, 3 * mm if kind in {"title", "heading"} else 1.5 * mm))
    document.build(story)
    return output.getvalue()


def export_bytes(version: ResearchBriefVersion, export_format: BriefExportFormat) -> tuple[bytes, str, str]:
    if export_format is BriefExportFormat.MARKDOWN:
        return render_markdown(version).encode("utf-8"), "text/markdown; charset=utf-8", "md"
    if export_format is BriefExportFormat.PDF:
        return render_pdf(version), "application/pdf", "pdf"
    return (
        render_docx(version),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    )
